"""DOCX解析器"""
import re
import os
from typing import List, Dict, Any, Optional
from loguru import logger
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.table import Table
from ..utils.config import Config

class DOCXParser:
    """DOCX解析器"""
    
    def __init__(self):
        self.chunk_size = Config.DOCX_CHUNK_SIZE
        self.chunk_overlap = Config.DOCX_CHUNK_OVERLAP
    
    def parse_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        解析DOCX文件
        
        Args:
            docx_path: DOCX文件路径
            
        Returns:
            解析后的文本块列表
        """
        logger.info(f"开始解析DOCX: {docx_path}")
        
        try:
            # 打开DOCX文件
            doc = Document(docx_path)
            
            chunks = []
            current_chapter = ""
            current_section = ""
            current_paragraph_num = 0
            
            # 解析文档内容
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    paragraph = Paragraph(element, doc)
                    current_paragraph_num += 1
                    
                    text = paragraph.text.strip()
                    if not text:
                        continue
                    
                    # 识别章节标题
                    chapter_info = self._extract_chapter_info(text, current_paragraph_num)
                    if chapter_info:
                        current_chapter = chapter_info["title"]
                        current_section = chapter_info.get("section", "")
                    
                    # 清理文本
                    cleaned_text = self._clean_text(text)
                    
                    # 分块处理
                    paragraph_chunks = self._split_text_into_chunks(
                        cleaned_text, 
                        current_paragraph_num,
                        current_chapter,
                        current_section
                    )
                    
                    chunks.extend(paragraph_chunks)
                    
                elif element.tag.endswith('tbl'):  # 表格
                    table = Table(element, doc)
                    current_paragraph_num += 1
                    
                    table_text = self._extract_table_text(table)
                    if table_text:
                        # 为表格创建特殊块
                        table_chunk = self._create_chunk(
                            table_text, 
                            current_paragraph_num,
                            current_chapter,
                            current_section,
                            chunk_type="table"
                        )
                        chunks.append(table_chunk)
            
            logger.info(f"✓ DOCX解析完成: {len(chunks)} 个文本块")
            return chunks
            
        except Exception as e:
            logger.error(f"✗ DOCX解析失败: {e}")
            raise
    
    def _extract_chapter_info(self, text: str, paragraph_num: int) -> Optional[Dict[str, str]]:
        """提取章节信息"""
        # 章节标题模式
        chapter_patterns = [
            r'第[一二三四五六七八九十\d]+章\s*([^\n]+)',
            r'第[一二三四五六七八九十\d]+节\s*([^\n]+)',
            r'^\s*(\d+\.\d*)\s+([^\n]+)',
            r'^\s*(\d+)\s+([^\n]+)',
            r'^\s*(\d+\.\d+\.\d*)\s+([^\n]+)',  # 子章节
            r'^\s*([一二三四五六七八九十]+)\s+([^\n]+)',  # 中文数字章节
        ]
        
        for pattern in chapter_patterns:
            match = re.search(pattern, text)
            if match:
                if len(match.groups()) == 1:
                    return {"title": match.group(1).strip()}
                else:
                    return {
                        "title": match.group(2).strip(),
                        "section": match.group(1).strip()
                    }
        
        return None
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除可能的页眉页脚
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 跳过可能的页眉页脚
            if self._is_header_footer(line):
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_header_footer(self, line: str) -> bool:
        """判断是否为页眉页脚"""
        # 页眉页脚特征
        header_footer_patterns = [
            r'^\d+$',  # 纯数字（页码）
            r'^第\s*\d+\s*页$',  # 第X页
            r'^\d+\s*/\s*\d+$',  # 页码格式
            r'^第[一二三四五六七八九十\d]+章',  # 重复的章节标题
        ]
        
        for pattern in header_footer_patterns:
            if re.match(pattern, line):
                return True
        
        # 过短的行可能是页眉页脚
        if len(line) < 3:
            return True
        
        return False
    
    def _extract_table_text(self, table: Table) -> str:
        """提取表格文本"""
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                table_text.append(" | ".join(row_text))
        
        return "\n".join(table_text)
    
    def _split_text_into_chunks(self, 
                               text: str, 
                               paragraph_num: int,
                               chapter: str,
                               section: str) -> List[Dict[str, Any]]:
        """将文本分割成块"""
        chunks = []
        
        # 如果文本长度小于chunk_size，直接作为一个块
        if len(text) <= self.chunk_size:
            chunk = self._create_chunk(
                text, 
                paragraph_num, 
                chapter, 
                section
            )
            chunks.append(chunk)
            return chunks
        
        # 按句子分割
        sentences = re.split(r'[。！？]', text)
        current_chunk = ""
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # 添加句号
            if not sentence.endswith(('。', '！', '？')):
                sentence += "。"
            
            # 如果当前块加上新句子会超过大小限制
            if current_length + len(sentence) > self.chunk_size and current_chunk:
                # 保存当前块
                chunk = self._create_chunk(
                    current_chunk, 
                    paragraph_num, 
                    chapter, 
                    section
                )
                chunks.append(chunk)
                
                # 开始新块（保留重叠部分）
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + sentence
                current_length = len(current_chunk)
            else:
                # 添加到当前块
                if current_chunk:
                    current_chunk += sentence
                else:
                    current_chunk = sentence
                current_length = len(current_chunk)
        
        # 保存最后一个块
        if current_chunk:
            chunk = self._create_chunk(
                current_chunk, 
                paragraph_num, 
                chapter, 
                section
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(self, 
                     content: str, 
                     paragraph_num: int,
                     chapter: str,
                     section: str,
                     chunk_type: str = "text") -> Dict[str, Any]:
        """创建文本块"""
        return {
            "content": content.strip(),
            "metadata": {
                "paragraph": paragraph_num,
                "chapter": chapter or f"第{paragraph_num}段",
                "section": section or "",
                "chunk_type": chunk_type
            }
        }
    
    def _get_overlap_text(self, text: str) -> str:
        """获取重叠文本"""
        if len(text) <= self.chunk_overlap:
            return text
        
        # 从末尾取重叠部分
        overlap = text[-self.chunk_overlap:]
        
        # 尝试在句子边界分割
        sentences = re.split(r'[。！？]', overlap)
        if len(sentences) > 1:
            # 保留最后一个完整句子
            return sentences[-1] + "。"
        
        return overlap
    
    def extract_metadata(self, docx_path: str) -> Dict[str, Any]:
        """提取DOCX元数据"""
        try:
            doc = Document(docx_path)
            
            # 计算段落数
            paragraph_count = len(doc.paragraphs)
            
            # 计算表格数
            table_count = len(doc.tables)
            
            result = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "creator": doc.core_properties.creator or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
                "paragraph_count": paragraph_count,
                "table_count": table_count
            }
            
            return result
            
        except Exception as e:
            logger.error(f"提取DOCX元数据失败: {e}")
            return {}
    
    def get_paragraph_text(self, docx_path: str, paragraph_num: int) -> str:
        """获取指定段落的文本"""
        try:
            doc = Document(docx_path)
            if paragraph_num < 1 or paragraph_num > len(doc.paragraphs):
                return ""
            
            paragraph = doc.paragraphs[paragraph_num - 1]
            text = paragraph.text
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"获取段落文本失败: {e}")
            return ""
    
    def batch_parse(self, 
                    metadata: Dict[str, Any],
                    book_ids: Optional[List[str]] = None,
                    versions: Optional[List[str]] = None) -> Dict[str, List[Dict[str, Any]]]:
        """批量解析DOCX文件"""
        logger.info("开始批量解析DOCX文件")
        
        all_chunks = {}
        
        for book in metadata["books"]:
            if book_ids and book["id"] not in book_ids:
                continue
            book_id = book["id"]
            book_name = book["name"]
            
            for version_info in book["versions"]:
                version = version_info["version"]
                if versions and version not in versions:
                    continue
                filename = version_info["filename"]
                
                # 只处理docx文件
                if not filename.endswith('.docx'):
                    continue
                
                # 构建collection名称
                collection_name = f"{book_id}_v{version}_docx"
                
                # 构建DOCX文件路径
                docx_path = os.path.join(Config.RAW_PDFS_DIR, filename)
                
                if not os.path.exists(docx_path):
                    logger.warning(f"DOCX文件不存在: {docx_path}")
                    continue
                
                logger.info(f"解析 {book_name} 第{version}版DOCX: {filename}")
                
                try:
                    # 解析DOCX
                    chunks = self.parse_docx(docx_path)
                    
                    # 为每个chunk添加书籍和版本信息
                    for chunk in chunks:
                        chunk["metadata"]["book_id"] = book_id
                        chunk["metadata"]["book_name"] = book_name
                        chunk["metadata"]["version"] = version
                        chunk["metadata"]["filename"] = filename
                        chunk["metadata"]["file_type"] = "docx"
                    
                    all_chunks[collection_name] = chunks
                    logger.info(f"✓ {book_name} 第{version}版DOCX解析完成: {len(chunks)} 个文本块")
                    
                except Exception as e:
                    logger.error(f"✗ 解析 {book_name} 第{version}版DOCX失败: {e}")
                    continue
        
        logger.info(f"批量解析完成，共生成 {len(all_chunks)} 个collections")
        return all_chunks

# 测试代码
if __name__ == "__main__":
    parser = DOCXParser()
    
    # 测试DOCX解析
    docx_path = "test.docx"  # 替换为实际的DOCX路径
    
    if os.path.exists(docx_path):
        chunks = parser.parse_docx(docx_path)
        print(f"解析完成，共 {len(chunks)} 个文本块")
        
        # 显示前几个块
        for i, chunk in enumerate(chunks[:3]):
            print(f"\n块 {i+1}:")
            print(f"段落: {chunk['metadata']['paragraph']}")
            print(f"章节: {chunk['metadata']['chapter']}")
            print(f"内容: {chunk['content'][:100]}...")
    else:
        print("DOCX文件不存在，请提供有效的DOCX路径")
